from docx import Document
import os

def generar_solicitud(nombre, fecha_inicio, fecha_fin, id_solicitud):
    documento = Document()

    documento.add_heading("SOLICITUD DE VACACIONES", level=1)

    documento.add_paragraph(f"Servidor: {nombre}")
    documento.add_paragraph(f"Fecha inicio: {fecha_inicio}")
    documento.add_paragraph(f"Fecha fin: {fecha_fin}")

    documento.add_paragraph(
        "Solicito la programación de mis vacaciones en las fechas indicadas."
    )

    if not os.path.exists("documentos"):
        os.makedirs("documentos")

    nombre_archivo = f"solicitud_{id_solicitud}.docx"
    ruta = os.path.join("documentos", nombre_archivo)

    documento.save(ruta)

    return nombre_archivo